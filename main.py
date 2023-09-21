from selenium import webdriver
import selenium.webdriver
import selenium.webdriver.common.keys
import selenium
import time
import os
import sys
import openai
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.action_chains import ActionChains
from bs4 import BeautifulSoup
import undetected_chromedriver as uc
import pyautogui
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

with open("apikey", "r") as file:
    apikey = file.read()

def generate_resume(i):
    openai.api_key = apikey
    response = openai.Completion.create(engine="text-davinci-003", prompt=f"Generate a cover letter for this description and company:{i} and make sure you include how much im willing to learn and even though I dont have any experience I can learn very quickly and that I will work for minimum wage and that I really just want to get experience and learning also only output the cover letter and nothing else", max_tokens=2000)
    text = response.choices[0].text.strip()
    print(text)
    doc = Document('Main Resume 2.0.docx')

    # Let's say you want to insert after the second paragraph
    target_paragraph = doc.paragraphs[1]

    insert_paragraph_after(target_paragraph, text=text)

    doc.save('output/ResumeGo.docx')


def linkdin():
    options = uc.ChromeOptions()
    chromeexecutablepath = "/opt/google/chrome/google-chrome"
    options.add_argument("user-data-dir=/home/dillon/.config/google-chrome/'profile 1'")
    #options.add_argument('--headless')
    options.add_argument('--disable-gpu')

    driver = uc.Chrome(chrome_options= options, executable_path=chromeexecutablepath)
    # Navigate to LinkedIn
    driver.get("https://www.linkedin.com")
    wait = WebDriverWait(driver, 10)

    with open("private", "r") as f:
        private = f.read()
        private = private.split("\n")

    email = wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="session_key"]')))
    email.send_keys(private[0])

    password = wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="session_password"]')))
    password.send_keys(private[1])

    enter = wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="main-content"]/section[1]/div/div/form/div[2]/button')))

    enter.click()

    driver.get("https://www.linkedin.com/jobs/search/?currentJobId=3661116542&distance=25&f_AL=true&f_E=2&geoId=103736294&keywords=Junior%20developer&origin=JOB_SEARCH_PAGE_JOB_FILTER&sortBy=R")

    time.sleep(4)



    jobs = driver.find_elements(By.CSS_SELECTOR, ".job-card-container")
    driver.execute_script("arguments[0].scrollIntoView();", jobs[-1])
    jobs = driver.find_elements(By.CSS_SELECTOR, ".job-card-container")
    driver.execute_script("arguments[0].scrollIntoView();", jobs[-1])
    jobs = driver.find_elements(By.CSS_SELECTOR, ".job-card-container")
    driver.execute_script("arguments[0].scrollIntoView();", jobs[-1])
    jobs = driver.find_elements(By.CSS_SELECTOR, ".job-card-container")
    driver.execute_script("arguments[0].scrollIntoView();", jobs[-1])
    jobs = driver.find_elements(By.CSS_SELECTOR, ".job-card-container")
    driver.execute_script("arguments[0].scrollIntoView();", jobs[-1])
    jobs = driver.find_elements(By.CSS_SELECTOR, ".job-card-container")
    print(jobs)

    for job in jobs:
        job.click()
        span = driver.find_element(By.XPATH, '//*[@id="job-details"]/span')



    time.sleep(100)
    driver.quit()

def test():
    options = uc.ChromeOptions()
    chromeexecutablepath = "/opt/google/chrome/google-chrome"
    options.add_argument("user-data-dir=/home/dillon/.config/google-chrome/'profile 1'")
    # options.add_argument('--headless')
    options.add_argument('--disable-gpu')

    driver = uc.Chrome(chrome_options=options, executable_path=chromeexecutablepath)
    # Navigate to LinkedIn
    driver.get("https://www.linkedin.com")
    time.sleep(200)
    driver.quit()
if __name__ == '__main__':
    generate_resume("""About the job
A wide array of extremely challenging C++ development tasks
An international team of brilliant minds
A competitive salary from the start and a raise to EUR 130,000 annually after only one year
A working environment that makes this team stay and grow
Enough time to make sure that every detail of your solution is perfect
A flat organization and plenty of room for your ideas
No scheduled meetings
Family-friendly working hours, no deadlines, no overtime
Support for relocation to Berlin or remote work opportunity!


About this job

We are looking for smart, creative developers with a solid theoretical background. You should be able to look at a problem from the user's perspective, discuss abstract concepts with fellow developers, as well as produce an elegant implementation. Developers we have hired in the past mostly hold an exceptional master's degree in computer science or even a doctorate.


At think-cell, we like flat hierarchies. You will work largely independently and will be responsible for the whole range of activities when implementing a new feature. We expect each of our developers to do architecture, design, implementation and bug fixing, rather than splitting these activities between several people. We thus minimize communication losses and put everyone in control of their own work. Your ideas are welcome, even if they mean that we have to change a lot of code to make things better.


We have published several scientific articles in the areas of Artificial Intelligence and Computer Graphics and we will encourage you to do the same. We sponsor visits to conferences and have close connections to universities and research institutes in the U.S. and Germany.


think-cell encourages a healthy work-life balance. We do not work at night or on weekends, and support our staff's families with a full-time company nanny. She is available for free when children are sick, or if you just feel like spending an evening out.


We pay very competitive salaries, and offer our developers EUR 130,000 annually following one year of employment. If necessary, we will go out of our way to help you relocate to Berlin, and will do what we can to help you acquire a work permit.


Remote work is now also possible!



About our software

Since 2002, think-cell produces graphics software that performs most of the painstaking work of creating data-driven slides for professional Excel and PowerPoint users. PowerPoint slide creation is one of the most popular things professionals use a computer for. Thus, it is rather surprising that while intelligent software has revolutionized many things we frequently do, such as web search in our browser, or telephone speech recognition in a call center, office productivity software has not changed much over the past decade or so.


think-cell is out to change this. We stand out from the crowd of other presentation software because we are willing to do the leg work of developing sophisticated algorithms and refining our user interface, which makes working with our software so satisfying for our users. Here are some highlights of what we have done.


Algorithmic highlights

We developed a new algorithm for automatic point cloud labeling that allows labels to be positioned away from the actual points.
We developed a new algorithm for automatic column chart labeling.
We are working with John Forrest – author of the linear solver CLP – to make his simplex code faster on our kind of problems.
We developed quite a few generic data structures that are not in C++ or Boost, for example partitions.


Hacking highlights

To do things that are not possible via the documented Microsoft Office API, we do lots of reverse engineering with the disassembler IDA from Hex-Rays.
We wrote probably the best function hooking engine out there. On each start of our software, we patch the Microsoft Office executables in memory. We search for small chunks of assembly code rather than hard-coding patch addresses to be robust against minor code modifications.
We redirect PowerPoint's and Excel's window contents into offscreen buffers and use Direct3D 9.0 to render our user interface on top.


Other highlights

We fund the working group for programming languages of the German Institute for Standardization (DIN). Some of our employees are members of this committee and vote in the international standardization process of ISO/IEC C++.
We closely track the latest versions of our compilers, Visual C++ and Xcode, so we can always use the latest C++ standard features as soon as they become available.
We use Boost throughout our code, e.g., Boost.Spirit for most of our parsing needs.
We have our own range library, in the same spirit as Boost.Range or Eric Niebler’s range-v3, but going further, for example, by unifying internal and external iteration. We gave a talk about it, and most of the code is public.
We have our own reference-counting and persistence libraries to save and restore whole object trees.
We wrote a parser and writer for the Excel .xls format.
We have an extensive bug reporting infrastructure. Assertions and error checks stay in the release code, and our software automatically reports bugs to our server. The server analyzes the bug, categorizes it and files it in a database that all developers can access. If an update fixes the bug, the user can download the update directly from a bug response web page.


Interested?

If you are the one missing from our team then please send us your electronic application via hr@think-cell.com


 We will only accept applications that contain a CV.

Pay found in job post
Retrieved from the description.

Is this accurate?

Yes
/

No
Base salary
$130,000/yr (from job description)
""")